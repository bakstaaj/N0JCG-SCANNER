#!/usr/bin/env python3
"""Build the branded N0JCG Scanner user manual from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "USER_MANUAL.md"
LOGO = ROOT / "web" / "assets" / "brand" / "N0JCG_Header_Dark_Approved.png"
OUTPUT = ROOT / "docs" / "publications" / "N0JCG_Scanner_User_Manual.docx"

NAVY = "0A1F44"
BLUE = "1565C0"
CYAN = "00B8D9"
SLATE = "2B3440"
MIST = "F4F7FA"
WHITE = "FFFFFF"
TEXT = "15202B"
MUTED = "6B7785"
WARNING = "B25E00"
DANGER = "C62828"

BODY_FONT = "Arial"
DISPLAY_FONT = "Arial"
MONO_FONT = "Consolas"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name=BODY_FONT, size=None, color=TEXT, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, *, left=None, bottom=None, size=16, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge, color in (("left", left), ("bottom", bottom)):
        if not color:
            continue
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), str(space))
        node.set(qn("w:color"), color)
        p_bdr.append(node)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=8.5, color=MUTED)


def normalize_public_text(text: str) -> str:
    replacements = {
        "PI Scanner": "N0JCG Scanner",
        "Start Scanning + Audio": "Start scanning + audio",
        "Radio Setup": "Radio setup",
        "Logs / Details": "Logs and details",
        "Skip 10 Min": "Skip 10 min",
        "Block Channel": "Block channel",
        "Clear Lock": "Clear lock",
        "Clear Blocks": "Clear blocks",
        "Full UI": "Full dashboard",
        "**Online:**": "**Connected:**",
        "**Online**": "**Connected**",
        "shows Online": "shows Connected",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, size=10.5, color=TEXT, bold=False) -> None:
    text = normalize_public_text(text)
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, size=max(8.5, size - 0.5), color=NAVY, bold=bold)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF5FC")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, bold=bold, italic=True)
        else:
            label = re.match(r"\[([^]]+)\]\(([^)]+)\)", token).group(1)
            run = paragraph.add_run(label)
            set_run_font(run, size=size, color=BLUE, bold=bold)
            run.underline = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color, bold=bold)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.keep_together = True

    for name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 18, 8),
        ("Heading 2", 14, BLUE, 14, 6),
        ("Heading 3", 11.5, SLATE, 10, 4),
    ):
        style = styles[name]
        style.font.name = DISPLAY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_together = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Leave room for the branded running header even when a paragraph or table
    # continues automatically onto a new page in LibreOffice.
    section.top_margin = Inches(1.08)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)
    # One page style keeps LibreOffice pagination consistent for automatic
    # paragraph and table continuations. The cover already carries full brand
    # artwork, so the compact running band is acceptable there as well.
    section.different_first_page_header_footer = False


def add_running_furniture(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    # Keep the running header intentionally empty. LibreOffice can overlay a
    # populated Word header on paragraphs moved by automatic pagination. The
    # branded cover and every-page footer provide consistent identity without
    # risking clipped operating instructions.

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    set_paragraph_border(p, bottom=CYAN, size=6, space=4)
    left = p.add_run("N0JCG Open Radio Platform  |  N0JCG Scanner v3.0.0        ")
    set_run_font(left, size=8.5, color=MUTED)
    page = p.add_run("Page ")
    set_run_font(page, size=8.5, color=MUTED)
    add_page_field(p)


def add_cover(doc: Document) -> None:
    banner = doc.add_table(rows=1, cols=1)
    set_table_geometry(banner, [CONTENT_WIDTH_DXA], indent=0)
    set_cell_shading(banner.cell(0, 0), NAVY)
    set_cell_margins(banner.cell(0, 0), top=260, bottom=260, start=300, end=300)
    p = banner.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(4.55))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(46)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(10)
    set_run_font(kicker.add_run("OPERATOR HANDBOOK"), name=DISPLAY_FONT, size=10, color=CYAN, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("N0JCG Scanner"), name=DISPLAY_FONT, size=31, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_run_font(subtitle.add_run("Installation, radio setup, operation, and troubleshooting"), size=14, color=SLATE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_border(rule, bottom=CYAN, size=12, space=1)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.65)
    summary.paragraph_format.right_indent = Inches(0.65)
    summary.paragraph_format.space_after = Pt(28)
    add_inline(summary, "A complete guide to the split-host P25, VHF, and UHF scanning system, browser audio, radio profiles, and stable RTL-SDR serial assignments.", size=11.5)

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680], indent=0)
    labels = (("RELEASE", "3.0.0"), ("PUBLICATION", "August 2026"), ("RADIO PATHS", "P25 / VHF / UHF"), ("AUDIENCE", "Operators and maintainers"))
    for index, (label, value) in enumerate(labels):
        row, col = divmod(index, 2)
        cell = meta.cell(row, col)
        set_cell_shading(cell, MIST)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(label + "\n"), size=8, color=BLUE, bold=True)
        set_run_font(p.add_run(value), size=10.5, color=NAVY, bold=True)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_quick_reference(doc: Document, headings: list[str], intro: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    add_inline(p, "Contents", size=18, color=NAVY, bold=True)
    set_paragraph_border(p, bottom=CYAN, size=8, space=4)
    lead = doc.add_paragraph()
    add_inline(lead, "Use this guide from initial hardware setup through daily operation and maintenance.")
    for heading in headings:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, re.sub(r"^\d+\.\s*", "", normalize_public_text(heading)))
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(12)
    shade_paragraph(callout, "E8FAFD")
    set_paragraph_border(callout, left=CYAN, size=18, space=8)
    add_inline(
        callout,
        "**Receiver ownership is serial-first.** Use the private station role "
        "map for VHF and UHF. Linux device indexes are never permanent role "
        "assignments.",
    )
    overview = doc.add_paragraph()
    overview.paragraph_format.space_before = Pt(10)
    shade_paragraph(overview, MIST)
    set_paragraph_border(overview, left=BLUE, size=18, space=8)
    add_inline(overview, normalize_public_text(intro))


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        if not is_table_separator(lines[index]):
            rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    weights = []
    for col in range(columns):
        length = max(len(re.sub(r"[`*]", "", row[col])) if col < len(row) else 0 for row in rows)
        weights.append(max(12, min(length, 52)))
    total = sum(weights)
    widths = [round(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_geometry(table, table_widths(rows))
    for table_row in table.rows:
        tr_pr = table_row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
    for row_index, source_row in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
                add_inline(p, source_row[col_index] if col_index < len(source_row) else "", size=9, color=WHITE, bold=True)
            else:
                if row_index % 2 == 0:
                    set_cell_shading(cell, "F8FAFC")
                add_inline(p, source_row[col_index] if col_index < len(source_row) else "", size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    shade_paragraph(p, "EEF2F6")
    set_paragraph_border(p, left=BLUE, size=10, space=7)
    set_run_font(p.add_run("\n".join(code_lines).rstrip()), name=MONO_FONT, size=8.3, color=SLATE)


def add_heading(doc: Document, level: int, text: str) -> None:
    style_name = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 3")
    p = doc.add_paragraph(style=style_name)
    size = {2: 18, 3: 14, 4: 11.5}.get(level, 11.5)
    color = {2: NAVY, 3: BLUE, 4: SLATE}.get(level, SLATE)
    add_inline(p, normalize_public_text(text), size=size, color=color, bold=True)
    if level == 2:
        set_paragraph_border(p, bottom=CYAN, size=8, space=4)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    text = normalize_public_text(text.strip())
    lowered = text.lower()
    if lowered.startswith(("important:", "warning:", "note:")) or text.startswith("The two analog assignments are mandatory"):
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        shade_paragraph(p, "FFF8EC" if not lowered.startswith("warning:") else "FFF2F2")
        set_paragraph_border(p, left=WARNING if not lowered.startswith("warning:") else DANGER, size=18, space=8)
    add_inline(p, text)


def new_numbering_instance(doc: Document, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = numbering.xpath(f'./w:num[@w:numId="{style_num_id}"]')[0]
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    direct_num_id = OxmlElement("w:numId")
    direct_num_id.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(direct_num_id)


def add_manual_body(doc: Document, lines: list[str]) -> None:
    index = 0
    paragraph_lines: list[str] = []
    active_numbering_id = None

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_lines))
            paragraph_lines = []

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            flush_paragraph()
            active_numbering_id = None
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            active_numbering_id = None
            index += 1
            code = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            add_code_block(doc, code)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            active_numbering_id = None
            add_heading(doc, len(heading.group(1)), heading.group(2))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            active_numbering_id = None
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            text = (bullet or numbered).group(1)
            source_number = int(stripped.split(".", 1)[0]) if numbered else None
            if numbered and (active_numbering_id is None or source_number == 1):
                active_numbering_id = new_numbering_instance(doc, source_number)
            if bullet:
                active_numbering_id = None
            index += 1
            while index < len(lines):
                continuation = lines[index].strip()
                if not continuation or re.match(r"^[-*]\s+|^\d+\.\s+|^#{2,4}\s+|^```|^\|", continuation):
                    break
                text += " " + continuation
                index += 1
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            if numbered:
                apply_numbering(p, active_numbering_id)
            add_inline(p, text)
            continue
        active_numbering_id = None
        paragraph_lines.append(stripped)
        index += 1
    flush_paragraph()


def source_parts(markdown: str) -> tuple[str, list[str], list[str]]:
    lines = markdown.splitlines()
    intro = []
    headings = []
    body_start = None
    in_contents = False
    for index, line in enumerate(lines):
        if line.startswith("## Contents"):
            in_contents = True
            continue
        if line.startswith("## ") and not line.startswith("## Contents"):
            headings.append(line[3:].strip())
            if body_start is None:
                body_start = index
            in_contents = False
        elif (
            body_start is None
            and not in_contents
            and line.strip()
            and not line.startswith(("# ", "|"))
        ):
            intro.append(line.strip())
    if body_start is None:
        raise ValueError("No numbered manual sections found")
    return " ".join(intro), headings, lines[body_start:]


def build() -> Path:
    if not SOURCE.is_file() or not LOGO.is_file():
        raise FileNotFoundError("Manual source or approved brand logo is missing")
    intro, headings, body_lines = source_parts(SOURCE.read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    add_running_furniture(section)
    add_cover(doc)
    add_quick_reference(doc, headings, intro)
    add_manual_body(doc, body_lines)

    properties = doc.core_properties
    properties.title = "N0JCG Scanner User Manual"
    properties.subject = "Installation, radio setup, operation, and troubleshooting"
    properties.author = "N0JCG Open Radio Platform"
    properties.keywords = "N0JCG, scanner, RTL-SDR, P25, VHF, UHF, user manual"
    properties.comments = "Generated from docs/USER_MANUAL.md"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
